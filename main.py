# --- Imports ---
import os
import io
import re
import httpx
import math
import pandas as pd
from fastapi import FastAPI, Response, UploadFile, File, Query
from fastapi.responses import StreamingResponse, JSONResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
from dotenv import load_dotenv
from pymongo import MongoClient
from starlette.background import BackgroundTask
try:
    from bson.objectid import ObjectId
except Exception:
    ObjectId = None
from datetime import datetime, timedelta
# For docx and pdf generation
from docx import Document
from xml.sax.saxutils import escape
import uuid
import shutil
import zipfile
import tempfile
try:
    from docx2pdf import convert as docx2pdf_convert
except ImportError:
    docx2pdf_convert = None


# --- Configuration ---
load_dotenv()
BACKEND_URL = os.environ.get("BACKEND_URL", "http://localhost:8080")
DATABASE_URL = os.environ.get("DATABASE_URL")

# Mongo client
_mongo_client = None
def get_mongo_client():
    global _mongo_client
    if _mongo_client is None:
        if not DATABASE_URL:
            raise RuntimeError('DATABASE_URL not set')
        _mongo_client = MongoClient(DATABASE_URL, serverSelectionTimeoutMS=5000)
    return _mongo_client

def try_objectid(s):
    if ObjectId is None:
        return s
    try:
        return ObjectId(s)
    except Exception:
        return s

# --- App Initialization ---
app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- Routes ---

# --- API: Get all applications for a product ---
@app.get('/api/applications')
def get_applications(productId: str = None):
    db = get_mongo_client().get_database('test')
    qry = {}
    if productId:
        qry['productId'] = try_objectid(productId)
    apps = list(db['policyapplications'].find(qry))
    data = [{
        '_id': str(a.get('_id')),
        'firstName': a.get('firstName', ''),
        'lastName': a.get('lastName', ''),
        'email': a.get('email', ''),
        'mobile': a.get('mobile', ''),
        'status': a.get('status', ''),
    } for a in apps]
    return JSONResponse({'data': data})

# --- API: Get all insured people for an application ---
@app.get('/api/insuredpeople')
def get_insured_people(applicationId: str = None):
    db = get_mongo_client().get_database('test')
    qry = {}
    if applicationId:
        qry['applicationId'] = try_objectid(applicationId)
    people = list(db['insuredpeople'].find(qry))
    data = [{
        '_id': str(p.get('_id')),
        'personIndex': p.get('personIndex'),
        'data': p.get('data', {}),
    } for p in people]
    return JSONResponse({'data': data})

# --- Member Details PDF Generation ---
@app.get("/api/member-details/{application_id}/pdf")
def generate_member_details_pdf(application_id: str, insuredId: str = None, format: str = "docx"):
    """
    Fetches data from DB, fills Member Details.docx, converts to PDF, returns PDF.
    """
    db = get_mongo_client().get_database('test')
    # Fetch application and insured people
    application = db['policyapplications'].find_one({"_id": try_objectid(application_id)})
    if not application:
        return JSONResponse({"error": "Application not found"}, status_code=404)
    insured_q = {"applicationId": try_objectid(application_id)}
    if insuredId:
        insured_q['_id'] = try_objectid(insuredId)
    insured_people = list(db['insuredpeople'].find(insured_q))

    # Fetch product and base products for policy details
    product = None
    policy_name = ""
    terms_list = []
    online_names = []
    online_desc = []
    offline_names = []
    offline_desc = []
    try:
        if application.get('productId'):
            product = db['products'].find_one({'_id': application.get('productId')})
        if product:
            policy_name = product.get('name', '') or ""
            base_ids = product.get('baseProduct') or []
            if base_ids:
                base_coll = db['baseproducts']
                for bid in base_ids:
                    lookup_id = try_objectid(bid)
                    bp = base_coll.find_one({'_id': lookup_id})
                    if not bp:
                        continue
                    if bp.get('termsConditions'):
                        terms_list.append(bp.get('termsConditions'))
                    name_val = bp.get('name') or "-"
                    desc_val = bp.get('detailDescription') or bp.get('termsConditions') or "-"
                    if bp.get('online') is True:
                        online_names.append(name_val)
                        online_desc.append(desc_val)
                    else:
                        offline_names.append(name_val)
                        offline_desc.append(desc_val)
    except Exception:
        policy_name = policy_name or ""
        terms_list = terms_list or []

    if not policy_name:
        policy_name = "Platinum Membership"
    # Single T&C (all base products share the same); take the first non-empty
    t_and_c_text = "\n".join(terms_list[:1]) if terms_list else ""
    if t_and_c_text:
        normalized = t_and_c_text.replace('\r\n', '\n').replace('\r', '\n')
        normalized = re.sub(r'\n{2,}', '\n', normalized)
        t_and_c_text = normalized.strip()
    t_and_c_value = t_and_c_text or "-"
    print(f"[T&C] application_id={application_id} terms={t_and_c_value}")

    # Online/offline base products, aligned line-wise (name line, description line)
    online_policy_name = "\n".join([x for x in online_names if x]) or "-"
    online_policy_desc = "\n".join([x for x in online_desc if x]) or "-"
    offline_policy_name = "\n".join([x for x in offline_names if x]) or "-"
    offline_policy_desc = "\n".join([x for x in offline_desc if x]) or "-"

    # Member/order details
    membership_num = str(application.get('_id', '')) or "-"
    purchase_date = None
    issue_date_str = "-"
    end_date_str = "-"
    customer_name = " ".join(filter(None, [application.get('firstName'), application.get('lastName')])) or "-"
    contact_number = application.get('mobile') or "-"
    dob = "-"
    gender = "-"
    nominee_name = "-"
    nominee_relation = "-"
    if insured_people:
        primary = insured_people[0]
        pdata = primary.get('data') or {}
        dob = pdata.get('dob') or "-"
        gender = pdata.get('gender') or "-"
        nominee_name = pdata.get('nomineeName') or "-"
        nominee_relation = pdata.get('nomineeRelation') or "-"

    # Payment/order amounts
    order_id = "-"
    net_price = "-"
    gst_val = "-"
    selling_price = "-"
    try:
        pay = db['paymentorders'].find_one({'applicationId': application.get('_id')}, sort=[('createdAt', -1)])
        if pay:
            order_id = str(pay.get('_id') or pay.get('razorpayPaymentId') or pay.get('razorpayPaymentLinkId') or '-')
            amount = pay.get('amount')
            purchase_date = pay.get('createdAt') or purchase_date
            if amount is not None:
                try:
                    selling_price_num = float(amount)
                    gst_num = round(selling_price_num * 0.18 / 1.18)
                    net_num = selling_price_num - gst_num
                    selling_price = f"{selling_price_num:.2f}"
                    gst_val = f"{gst_num:.2f}"
                    net_price = f"{net_num:.2f}"
                except Exception:
                    pass
    except Exception:
        pass

    if purchase_date is None:
        purchase_date = application.get('createdAt')
    if isinstance(purchase_date, datetime):
        issue_date_str = purchase_date.date().isoformat()
        end_date_str = (purchase_date + timedelta(days=365)).date().isoformat()
    elif purchase_date:
        issue_date_str = str(purchase_date)

    insurer = "-"
    policy_number = "-"
    # Prepare docx template
    template_path = os.path.join(os.path.dirname(__file__), "Member Details.docx")
    if not os.path.exists(template_path):
        return JSONResponse({"error": "Template DOCX not found"}, status_code=500)
    doc = Document(template_path)
    # Replace placeholders in the docx
    placeholders = {
        "{{firstName}}": application.get("firstName") or "-",
        "{{lastName}}": application.get("lastName") or "-",
        "{{email}}": application.get("email") or "-",
        "{{mobile}}": contact_number,
        "{{#policy_name}}": policy_name,
        "{{policy_name}}": policy_name,
        "{{#t_and_c}}": t_and_c_value,
        "{{t_and_c}}": t_and_c_value,
        "{{policy_name}} (this is where t and c are going to be there)": t_and_c_value,
        "{{customer_name}}": customer_name,
        "{{full_name}}": customer_name,
        "{{first_name}}": application.get("firstName") or "-",
        "{{last_name}}": application.get("lastName") or "-",
        "{{phone_number}}": contact_number,
        "{{contact_number}}": contact_number,
        "{{address}}": application.get("address") or "-",
        "{{city}}": application.get("city") or "-",
        "{{state}}": application.get("state") or "-",
        "{{country}}": application.get("country") or "-",
        "{{pincode}}": application.get("pincode") or "-",
        "{{gender}}": gender,
        "{{dob}}": dob,
        "{{nominee_name}}": nominee_name,
        "{{nominee_relation}}": nominee_relation,
        "{{member_num}}": membership_num,
        "{{issue_date}}": issue_date_str,
        "{{ issue_date }}": issue_date_str,
        "{{end_date}}": end_date_str,
        "{{ end_date }}": end_date_str,
        "{{order_id}}": order_id,
        "{{net_price}}": net_price,
        "{{gst}}": gst_val,
        "{{selling_price}}": selling_price,
        "{{insurer}}": insurer,
        "{{policyNumber}}": policy_number,
        "{{policy_number}}": policy_number,
        "{{/policy_name}}": "",
        "{{#online_policy_name}}": online_policy_name,
        "{{online_policy_name}}": online_policy_name,
        "{{#online_policy_description}}": online_policy_desc,
        "{{online_policy_description}}": online_policy_desc,
        "{{#offline_policy_name}}": offline_policy_name,
        "{{offline_policy_name}}": offline_policy_name,
        "{{#offline_policy_description}}": offline_policy_desc,
        "{{offline_policy_description}}": offline_policy_desc,
    }

    def replace_in_paragraphs(paragraphs, mapping):
        for para in paragraphs:
            for ph, val in mapping.items():
                if ph in para.text:
                    for run in para.runs:
                        if ph in run.text:
                            run.text = run.text.replace(ph, str(val))
                    para.text = para.text.replace(ph, str(val))

    def replace_in_tables(tables, mapping):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_in_paragraphs(cell.paragraphs, mapping)

    def replace_in_headers_footers(document, mapping):
        for section in document.sections:
            replace_in_paragraphs(section.header.paragraphs, mapping)
            replace_in_tables(section.header.tables, mapping)
            replace_in_paragraphs(section.footer.paragraphs, mapping)
            replace_in_tables(section.footer.tables, mapping)

    replace_in_paragraphs(doc.paragraphs, placeholders)
    replace_in_tables(doc.tables, placeholders)
    replace_in_headers_footers(doc, placeholders)

    # Fallback: replace a literal title if present
    if policy_name:
        replace_in_paragraphs(doc.paragraphs, {"Platinum Membership": policy_name})
        replace_in_tables(doc.tables, {"Platinum Membership": policy_name})

    # TODO: If the template has a table for insured people, populate it here using insured_people.
    # Save filled docx with applicant name in filename to aid debugging
    base_name = "_".join([application.get("firstName", ""), application.get("lastName", "")]).strip("_")
    safe_name = re.sub(r"[^A-Za-z0-9_-]+", "", base_name) or "member"
    temp_docx = f"/tmp/member_details_{safe_name}_{uuid.uuid4().hex}.docx"
    doc.save(temp_docx)

    # Raw XML replace to catch placeholders inside shapes/textboxes
    try:
        def raw_replace(docx_path, mapping):
            with tempfile.TemporaryDirectory() as td:
                with zipfile.ZipFile(docx_path, 'r') as zin:
                    zin.extractall(td)
                # Replace in all XML parts
                for root, _, files in os.walk(td):
                    for fn in files:
                        if fn.endswith('.xml'):
                            p = os.path.join(root, fn)
                            try:
                                with open(p, 'r', encoding='utf-8') as f:
                                    content = f.read()
                                for ph, val in mapping.items():
                                    safe_val = escape(str(val))
                                    content = content.replace(ph, safe_val)
                                    if ph.startswith('{{') and ph.endswith('}}'):
                                        core = ph[2:-2].strip()
                                        if core:
                                            pattern = (
                                                r"\{\{(?:[\s\u2028\u2029\u00A0]|<[^>]+>|&[a-zA-Z]+;|&#\d+;)*"
                                                + re.escape(core)
                                                + r"(?:[\s\u2028\u2029\u00A0]|<[^>]+>|&[a-zA-Z]+;|&#\d+;)*\}\}"
                                            )
                                            content = re.sub(pattern, safe_val, content)
                                with open(p, 'w', encoding='utf-8') as f:
                                    f.write(content)
                            except Exception:
                                pass
                # Repack
                with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zout:
                    for root, _, files in os.walk(td):
                        for fn in files:
                            fp = os.path.join(root, fn)
                            arcname = os.path.relpath(fp, td)
                            zout.write(fp, arcname)
        raw_mapping = dict(placeholders)
        raw_mapping["Platinum Membership"] = policy_name or raw_mapping.get("Platinum Membership", "")
        raw_replace(temp_docx, raw_mapping)
    except Exception:
        pass
    # Decide output format
    want_pdf = str(format).lower() == "pdf"
    temp_pdf = temp_docx.replace(".docx", ".pdf")
    pdf_ready = False

    if want_pdf and docx2pdf_convert:
        try:
            docx2pdf_convert(temp_docx, temp_pdf)
            pdf_ready = os.path.exists(temp_pdf) and os.path.getsize(temp_pdf) > 0
        except Exception:
            pdf_ready = False

    if want_pdf and pdf_ready:
        def cleanup_files():
            for path in (temp_docx, temp_pdf):
                try:
                    os.remove(path)
                except Exception:
                    pass
        return FileResponse(
            temp_pdf,
            media_type="application/pdf",
            filename="MemberDetails.pdf",
            background=BackgroundTask(cleanup_files),
            headers={"Content-Disposition": 'attachment; filename="MemberDetails.pdf"'}
        )

    # Fallback to DOCX (either requested explicitly or PDF unavailable)
    def cleanup_docx():
        try:
            os.remove(temp_docx)
        except Exception:
            pass
        try:
            if os.path.exists(temp_pdf):
                os.remove(temp_pdf)
        except Exception:
            pass
    return FileResponse(
        temp_docx,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="MemberDetails.docx",
        background=BackgroundTask(cleanup_docx),
        headers={"Content-Disposition": 'attachment; filename="MemberDetails.docx"'}
    )


# --- Member Details DOCX direct endpoint ---
@app.get("/api/member-details/{application_id}/docx")
def generate_member_details_docx(application_id: str, insuredId: str = None):
    return generate_member_details_pdf(application_id, insuredId=insuredId, format="docx")

# --- Routes ---
@app.get("/")
def serve_index():
    index_path = os.path.join(os.path.dirname(__file__), "index.html")
    return FileResponse(index_path, media_type="text/html")

@app.get("/api/products/active")
def get_active_products():
    url = f"{BACKEND_URL}/api/products/active"
    try:
        with httpx.Client(timeout=10) as client:
            resp = client.get(url)
            return JSONResponse(resp.json(), status_code=resp.status_code)
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)


def _range_to_dates(range_key: str):
    now = pd.Timestamp.now()
    if range_key == '1D':
        start = now - pd.Timedelta(days=1)
    elif range_key == '1W':
        start = now - pd.Timedelta(weeks=1)
    elif range_key == '1M':
        start = now - pd.DateOffset(months=1)
    elif range_key == 'last_month':
        start = (now.replace(day=1) - pd.DateOffset(months=1)).normalize()
        end = now.replace(day=1) - pd.DateOffset(days=1)
        return pd.Timestamp(start.to_pydatetime()).to_pydatetime(), pd.Timestamp(end.to_pydatetime()).to_pydatetime()
    elif range_key == 'this_quarter':
        q = (now.month - 1) // 3 + 1
        start = pd.Timestamp(datetime(now.year, 3 * (q - 1) + 1, 1))
    elif range_key == 'last_quarter':
        q = (now.month - 1) // 3 + 1
        start = pd.Timestamp(datetime(now.year, 3 * (q - 2) + 1, 1))
    elif range_key == 'this_fy':
        fy_start_month = 4
        if now.month >= fy_start_month:
            start = pd.Timestamp(datetime(now.year, fy_start_month, 1))
        else:
            start = pd.Timestamp(datetime(now.year - 1, fy_start_month, 1))
    elif range_key == 'last_fy':
        fy_start_month = 4
        if now.month >= fy_start_month:
            start = pd.Timestamp(datetime(now.year - 1, fy_start_month, 1))
        else:
            start = pd.Timestamp(datetime(now.year - 2, fy_start_month, 1))
    else:
        start = pd.Timestamp.min
    end = pd.Timestamp.now()
    return pd.Timestamp(start.to_pydatetime()).to_pydatetime(), pd.Timestamp(end.to_pydatetime()).to_pydatetime()


@app.get('/api/meta/ranges')
def meta_ranges():
    return JSONResponse({
        'ranges': ['1D','1W','1M','last_month','this_quarter','last_quarter','this_fy','last_fy','all_time']
    })


@app.get('/api/policies')
def list_policies(q: str = None, page: int = 1, limit: int = 50):
    db = get_mongo_client().get_database('test')
    coll = db['products']
    qry = {}
    if q:
        qry['name'] = {'$regex': q, '$options': 'i'}
    total = coll.count_documents(qry)
    docs = list(coll.find(qry).skip((page-1)*limit).limit(limit))
    data = [{'policyId': str(d.get('_id')), 'policyName': d.get('name')} for d in docs]
    return JSONResponse({'data': data, 'pagination': {'page': page, 'limit': limit, 'total': total}})


@app.get('/api/agents/{agent_id}/dashboard/stats')
def agent_stats(agent_id: str, range: str = Query('1M')):
    db = get_mongo_client().get_database('test')
    start, end = _range_to_dates(range)
    # Handle agentId as ObjectId or string
    agent_q = {'$or': [{'agentId': agent_id}]}
    if ObjectId is not None:
        try:
            agent_oid = ObjectId(agent_id)
            agent_q['$or'].append({'agentId': agent_oid})
        except Exception:
            pass
    apps = list(db['policyapplications'].find({'$or': [{'agentId': agent_id}, {'agentId': try_objectid(agent_id)}]}))
    print(f"[DEBUG] Found {len(apps)} policyapplications for agentId {agent_id}")
    app_ids = [a.get('_id') for a in apps]

    # Map applicationId to productId
    appid_to_product = {str(a.get('_id')): str(a.get('productId')) if a.get('productId') else None for a in apps}

    # Issuances for this agent in range
    iss_q = {'applicationId': {'$in': app_ids}, 'createdAt': {'$gte': start, '$lte': end}} if app_ids else {'createdAt': {'$gte': start, '$lte': end}}
    print(f"[DEBUG] Issuance query: {iss_q}")
    issuances = list(db['policyissuances'].find(iss_q)) if app_ids else []
    print(f"[DEBUG] Found {len(issuances)} policyissuances for agentId {agent_id}")
    policies_issued = len(issuances)

    # Payments for this agent in range
    pay_q = {'applicationId': {'$in': app_ids}, 'createdAt': {'$gte': start, '$lte': end}} if app_ids else {'createdAt': {'$gte': start, '$lte': end}}
    print(f"[DEBUG] Payment query: {pay_q}")
    payments = list(db['paymentorders'].find(pay_q)) if app_ids else []
    print(f"[DEBUG] Found {len(payments)} paymentorders for agentId {agent_id}")
    # Only count successful payments for revenue
    ok_status = set(['paid','success','captured','completed','PAID','SUCCESS','CAPTURED','COMPLETED'])
    revenue = 0
    for p in payments:
        if str(p.get('status', '')).lower() in ok_status:
            try:
                amt = int(p.get('amount') or 0)
            except Exception:
                try:
                    amt = int(float(p.get('amount') or 0))
                except Exception:
                    amt = 0
            revenue += amt

    # Aggregate sales and revenue per product
    prod_count = {}
    for ins in issuances:
        # Try to get productId from issuance, else from application
        pid = ins.get('productId')
        if not pid:
            pid = appid_to_product.get(str(ins.get('applicationId')))
        pid = str(pid) if pid else 'unknown'
        entry = prod_count.setdefault(pid, {'sold': 0, 'revenue': 0})
        entry['sold'] += 1
    included_paymentorders = []
    for p in payments:
        if str(p.get('status', '')).lower() in ok_status:
            app_id = str(p.get('applicationId'))
            pid = appid_to_product.get(app_id) or 'unknown'
            pid = str(pid)
            try:
                amt = int(p.get('amount') or 0)
            except Exception:
                amt = 0
            print(f"[DEBUG] Summing paymentorder: _id={p.get('_id')} applicationId={app_id} amount={amt} status={p.get('status')} mappedProductId={pid}")
            included_paymentorders.append({'_id': str(p.get('_id')), 'amount': amt})
            prod_count.setdefault(pid, {'sold': 0, 'revenue': 0})
            prod_count[pid]['revenue'] += amt

    print(f"[DEBUG] Total revenue returned: {revenue}")
    print(f"[DEBUG] Paymentorders included in revenue: {[x['_id'] for x in included_paymentorders]}")

    total_revenue = sum(v['revenue'] for v in prod_count.values()) if prod_count else 0
    # Resolve product names
    prod_names = {}
    try:
        prod_coll = db['products']
        for pid in list(prod_count.keys()):
            if pid and pid != 'unknown':
                try:
                    lookup = try_objectid(pid)
                    doc = prod_coll.find_one({'_id': lookup})
                except Exception:
                    doc = prod_coll.find_one({'_id': pid})
                if doc and doc.get('name'):
                    prod_names[pid] = doc.get('name')
                else:
                    prod_names[pid] = pid
            else:
                prod_names[pid] = 'Unknown'
    except Exception:
        for pid in prod_count.keys():
            prod_names[pid] = pid

    top_list = []
    for pid, v in prod_count.items():
        top_list.append({
            'policyId': pid,
            'policyName': prod_names.get(pid, pid),
            'sold': v['sold'],
            'revenue': v['revenue'],
            'contribution': int((v['revenue']/total_revenue*100) if total_revenue else 0)
        })
    top_list = sorted(top_list, key=lambda x: x['sold'], reverse=True)
    best = top_list[0] if top_list else None

    # Multiply all revenue values by 100 for paise convention
    revenue_paise = revenue * 100
    for item in top_list:
        item['revenue'] = item['revenue'] * 100
    if best:
        best['revenue'] = best['revenue'] * 100
    return JSONResponse({
        'policiesIssued': policies_issued,
        'revenue': revenue_paise,
        'bestSellingPolicy': best,
        'topPolicies': top_list,
        'period': {'from': start.isoformat(), 'to': end.isoformat()}
    })


@app.get('/api/agents/{agent_id}/dashboard/chart')
def agent_chart(agent_id: str, start: str = None, end: str = None, interval: str = 'day', policyId: str = None):
    db = get_mongo_client().get_database('test')
    if start:
        start_dt = datetime.fromisoformat(start)
    else:
        start_dt, _ = _range_to_dates('1M')
    if end:
        end_dt = datetime.fromisoformat(end)
    else:
        _, end_dt = _range_to_dates('1M')
    apps = list(db['policyapplications'].find({'agentId': agent_id}))
    app_ids = [a.get('_id') for a in apps]
    if not app_ids:
        return JSONResponse({'series': [], 'meta': {'interval': interval, 'points': 0}})
    q = {'applicationId': {'$in': app_ids}, 'createdAt': {'$gte': start_dt, '$lte': end_dt}}
    issuances = list(db['policyissuances'].find(q))
    rows = []
    for ins in issuances:
        dt = ins.get('createdAt')
        date_key = dt.date().isoformat()
        rows.append({'date': date_key, 'revenue': int(ins.get('amount') or 0), 'sales': 1, 'productId': str(ins.get('productId') or '')})
    df = pd.DataFrame(rows)
    if df.empty:
        return JSONResponse({'series': [], 'meta': {'interval': interval, 'points': 0}})
    grp = df.groupby('date').agg({'sales':'sum','revenue':'sum'}).reset_index()
    series = grp.to_dict(orient='records')
    return JSONResponse({'series': series, 'meta': {'interval': interval, 'points': len(series)}})


@app.get('/api/agents/{agent_id}/dashboard/policies')
def agent_policies(agent_id: str, page: int = 1, limit: int = 10, sortBy: str = 'sold', order: str = 'desc', range: str = '1M'):
    db = get_mongo_client().get_database('test')
    start, end = _range_to_dates(range)
    # Handle agentId as ObjectId or string
    apps = list(db['policyapplications'].find({'$or': [{'agentId': agent_id}, {'agentId': try_objectid(agent_id)}]}))
    print(f"[DEBUG] Found {len(apps)} policyapplications for agentId {agent_id}")
    app_ids = [a.get('_id') for a in apps]
    if not app_ids:
        print("[DEBUG] No applications found for agent, returning empty data.")
        return JSONResponse({'data': [], 'pagination': {'page': page, 'limit': limit, 'total': 0}})
    appid_to_product = {str(a.get('_id')): str(a.get('productId')) if a.get('productId') else None for a in apps}
    iss_q = {'applicationId': {'$in': app_ids}, 'createdAt': {'$gte': start, '$lte': end}}
    print(f"[DEBUG] Issuance query: {iss_q}")
    issuances = list(db['policyissuances'].find(iss_q))
    print(f"[DEBUG] Found {len(issuances)} policyissuances for agentId {agent_id}")
    perf = {}
    for ins in issuances:
        pid = ins.get('productId')
        if not pid:
            pid = appid_to_product.get(str(ins.get('applicationId')))
        pid = str(pid) if pid else 'unknown'
        entry = perf.setdefault(pid, {'sold': 0, 'revenue': 0})
        entry['sold'] += 1
    payments = list(db['paymentorders'].find({'applicationId': {'$in': app_ids}, 'createdAt': {'$gte': start, '$lte': end}}))
    ok_status = set(['paid','success','captured','completed','PAID','SUCCESS','CAPTURED','COMPLETED'])
    for p in payments:
        if str(p.get('status', '')).lower() in ok_status:
            pid = appid_to_product.get(str(p.get('applicationId'))) or 'unknown'
            pid = str(pid)
            try:
                amt = int(p.get('amount') or 0)
            except Exception:
                amt = 0
            perf.setdefault(pid, {'sold': 0, 'revenue': 0})
            perf[pid]['revenue'] += amt
    # Resolve product names
    prod_names = {}
    try:
        prod_coll = db['products']
        for pid in list(perf.keys()):
            if pid and pid != 'unknown':
                try:
                    lookup = try_objectid(pid)
                    doc = prod_coll.find_one({'_id': lookup})
                except Exception:
                    doc = prod_coll.find_one({'_id': pid})
                if doc and doc.get('name'):
                    prod_names[pid] = doc.get('name')
                else:
                    prod_names[pid] = pid
            else:
                prod_names[pid] = 'Unknown'
    except Exception:
        for pid in perf.keys():
            prod_names[pid] = pid
    total_revenue = sum(x['revenue'] for x in perf.values()) if perf else 0
    rows = []
    for pid, v in perf.items():
        rows.append({
            'policyId': pid,
            'policyName': prod_names.get(pid, pid),
            'icon': None,
            'sold': v['sold'],
            'revenue': v['revenue'] * 100,
            'contribution': int((v['revenue']/total_revenue*100) if total_revenue else 0)
        })
    if sortBy in ('sold','revenue','contribution'):
        rows = sorted(rows, key=lambda x: x.get(sortBy,0), reverse=(order=='desc'))
    total_items = len(rows)
    start_i = (page-1)*limit
    page_rows = rows[start_i:start_i+limit]
    return JSONResponse({'data': page_rows, 'pagination': {'page': page, 'limit': limit, 'total': total_items}})


@app.get('/api/agents/{agent_id}/sales')
def agent_sales(agent_id: str, page: int = 1, limit: int = 50):
    db = get_mongo_client().get_database('test')
    apps = list(db['policyapplications'].find({'agentId': agent_id}))
    app_ids = [a.get('_id') for a in apps]
    payments = list(db['paymentorders'].find({'applicationId': {'$in': app_ids}})) if app_ids else []
    rows = []
    for p in payments:
        rows.append({'id': str(p.get('_id')), 'applicationId': str(p.get('applicationId')), 'amount': p.get('amount'), 'status': p.get('status'), 'createdAt': p.get('createdAt')})
    total = len(rows)
    start_i = (page-1)*limit
    return JSONResponse({'data': rows[start_i:start_i+limit], 'pagination': {'page': page, 'limit': limit, 'total': total}})

@app.get("/excel/{product_id}")
def download_product_excel(product_id: str):
    url = f"{BACKEND_URL}/api/products/active"
    try:
        with httpx.Client(timeout=10) as client:
            resp = client.get(url)
            products = resp.json().get("data", [])
    except Exception as e:
        return JSONResponse({"error": f"Failed to fetch products: {e}"}, status_code=500)
    product = next((p for p in products if p["_id"] == product_id), None)
    if not product:
        return JSONResponse({"error": "Product not found"}, status_code=404)
    columns = [f["fieldId"]["fieldName"] for f in product["fields"]]
    validations = {}
    for idx, f in enumerate(product["fields"]):
        field_type = f["fieldId"].get("dataType", "")
        options = f["fieldId"].get("options", [])
        if field_type in ["dropdown", "radio", "checkbox"] and options:
            validations[idx] = options
    sample_row = []
    for idx, f in enumerate(product["fields"]):
        field_type = f["fieldId"].get("dataType", "")
        options = f["fieldId"].get("options", [])
        if field_type in ["dropdown", "radio", "checkbox"] and options:
            sample_row.append(options[0])
        elif field_type == "number":
            sample_row.append(12345)
        elif field_type == "adharCard":
            sample_row.append("123412341234")
        elif field_type == "panCard":
            sample_row.append("ABCDE1234F")
        elif field_type == "email":
            sample_row.append("sample@email.com")
        else:
            sample_row.append("Sample Value")
    df = pd.DataFrame([sample_row], columns=columns)
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        df.to_excel(writer, index=False)
        workbook = writer.book
        worksheet = writer.sheets["Sheet1"]
        for col_idx, opts in validations.items():
            col_letter = chr(ord('A') + col_idx)
            worksheet.data_validation(f'{col_letter}2:{col_letter}101', {
                'validate': 'list',
                'source': opts,
                'input_message': 'Select from dropdown',
            })
    output.seek(0)
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={
            "Content-Disposition": f"attachment; filename={product['name'].replace(' ', '_')}.xlsx"
        },
    )

@app.post("/validate-excel/{product_id}")
async def validate_excel(product_id: str, file: UploadFile = File(...), output: str = Query("json", enum=["json", "csv", "xlsx", "excel"])):
    url = f"{BACKEND_URL}/api/products/active"
    try:
        with httpx.Client(timeout=10) as client:
            resp = client.get(url)
            products = resp.json().get("data", [])
    except Exception as e:
        return JSONResponse({"error": f"Failed to fetch products: {e}"}, status_code=500)
    product = next((p for p in products if p["_id"] == product_id), None)
    if not product:
        return JSONResponse({"error": "Product not found"}, status_code=404)
    contents = await file.read()
    df = pd.read_excel(io.BytesIO(contents))
    field_rules = []
    for f in product["fields"]:
        field_type = f["fieldId"].get("dataType", "text")
        options = f["fieldId"].get("options", [])
        field_rules.append({
            "type": field_type,
            "options": options
        })
    errors = []
    for idx, row in df.iloc[1:].iterrows(): 
        for col_idx, value in enumerate(row):
            rule = field_rules[col_idx]
            t = rule["type"]
            opts = rule["options"]
            val = str(value).strip() if pd.notnull(value) else ""
            if t == "number":
                if val and not val.isdigit():
                    errors.append({"row": idx + 2, "column": col_idx + 1, "message": "Not a valid number", "value": val})
            elif t == "email":
                if val and not re.match(r"^[\w\.-]+@[\w\.-]+\.\w+$", val):
                    errors.append({"row": idx + 2, "column": col_idx + 1, "message": "Not a valid email", "value": val})
            elif t == "adharCard":
                val_no_space = val.replace(' ', '')
                if val and not re.match(r"^\d{12}$", val_no_space):
                    errors.append({"row": idx + 2, "column": col_idx + 1, "message": "Not a valid Aadhaar (12 digits, spaces ignored)", "value": val})
            elif t == "panCard":
                if val and not re.match(r"^[A-Z]{5}\d{4}[A-Z]$", val, re.I):
                    errors.append({"row": idx + 2, "column": col_idx + 1, "message": "Not a valid PAN (e.g. ABCDE1234F)", "value": val})
            elif t == "date":
                if val:
                    try:
                        pd.to_datetime(val)
                    except Exception:
                        errors.append({"row": idx + 2, "column": col_idx + 1, "message": "Not a valid date", "value": val})
            elif t in ["dropdown", "radio", "checkbox"]:
                if val and val not in opts:
                    errors.append({"row": idx + 2, "column": col_idx + 1, "message": "Value not in allowed options", "value": val, "allowedOptions": opts})
            # text: always valid
    if errors:
        if output in ("csv",):
            buffer = io.StringIO()
            pd.DataFrame(errors)[["row", "column", "message", "value"]].to_csv(buffer, index=False)
            buffer.seek(0)
            return StreamingResponse(
                io.BytesIO(buffer.getvalue().encode("utf-8")),
                media_type="text/csv",
                headers={"Content-Disposition": "attachment; filename=validation_errors.csv"},
            )
        if output in ("xlsx", "excel"):
            buffer = io.BytesIO()
            pd.DataFrame(errors)[["row", "column", "message", "value"]].to_excel(buffer, index=False)
            buffer.seek(0)
            return StreamingResponse(
                buffer,
                media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                headers={"Content-Disposition": "attachment; filename=validation_errors.xlsx"},
            )
        messages = [f"Row {e['row']}, Column {e['column']}: {e['message']}" for e in errors]
        return JSONResponse({"valid": False, "errors": messages})
    return JSONResponse({"valid": True, "message": "Excel is valid."})


